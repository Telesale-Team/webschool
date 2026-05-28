"""
Management command to fix empty choices for problem questions.

Root cause analysis:
- Questions with choices as OMML math: python-docx .text reads math as empty string.
  The original import also misread the "1. [MATH] 2. [MATH] 3. [MATH] 4. [MATH] 5. [MATH]"
  layout — the regex matched "2." and "4." as choice values for choice1 and choice3.
  So choice1="2." choice3="4." are WRONG data too and must be overwritten.

- Questions with choices in a nested table (data values): extract from nested table cells.
- Q87: choices are images -> placeholder text.
- Q156: choices are mixed text + 3 math blocks inline in a paragraph.
- Round 2 (20 questions): same MATH_5/MATH_PARA patterns plus plain-text paragraph parsing
  for questions where choices were spread across "N. value" inline or multi-line paragraphs.

Strategy map:
  MATH_5      -> math_blocks has exactly 5 items -> map directly to choices 1-5 (overwrite all)
  MATH_5+NESTED_LABELS -> nested table has only "1." "2." ... (no values),
                          math_blocks has 5 items -> use math_blocks (overwrite all)
  MATH_PARA   -> math_blocks > 5 (stem has math too), choices are per-paragraph
                 each "N. [MATH]" paragraph -> choices[N] = math_text
  NESTED_DATA -> nested table has real choice values -> extract from nested table
  IMAGE       -> insert placeholder for all 5 choices
  Q156        -> special: 3 math blocks in choice para, other 2 choices are plain text "2" and "5"
  PLAIN_INLINE -> choices inline in one paragraph "1.val  2.val  3.val  4.val  5.val"
  PLAIN_MULTILINE -> choices spread across multiple paragraphs
"""
import re
import docx
from django.core.management.base import BaseCommand
from questions.models import Question, Choice


# ─── XML helpers ────────────────────────────────────────────────────────────

def extract_math_blocks(cell_elem):
    """Return list of strings, one per oMath element in document order."""
    ns = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
    results = []
    for me in cell_elem.findall(f'.//{ns}oMath'):
        t_elems = me.findall(f'.//{ns}t')
        results.append(''.join(t.text or '' for t in t_elems).strip())
    return results


def extract_math_by_para(cell_elem):
    """
    Scan paragraphs for "N. [math]" pattern.
    Returns dict {choice_num (int): math_text (str)}.
    """
    ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    ns_m = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
    choices = {}
    for p in cell_elem.findall(f'.//{ns_w}p'):
        plain = ''.join(t.text or '' for t in p.findall(f'.//{ns_w}t')).strip()
        math_texts = []
        for me in p.findall(f'.//{ns_m}oMath'):
            math_texts.append(
                ''.join(t.text or '' for t in me.findall(f'.//{ns_m}t')).strip()
            )
        markers = re.findall(r'([1-5])\.', plain)
        for i, m in enumerate(markers):
            if i < len(math_texts):
                choices[int(m)] = math_texts[i]
    return choices


def nested_table_is_choice_table(cell):
    """
    Return True if nested table looks like a choice table with actual values.

    Supports two cell patterns:
    a) Alternating "N." | value cells (e.g. "1." | "76" | "2." | "113" ...)
       -> require that the value cell is NON-EMPTY and NOT another bare label
    b) Single combined cells "N. value" where value is non-empty
       (e.g. "1. 76", "2.  113", "3. 296")
       -> require the value part after "N." to be a non-label string

    Returns False for:
    - Tables with only bare labels like "1." "2." "3." "4." "5." (no values)
    - Data/frequency tables (headers like "อายุ", "จำนวน")
    """
    for nt in cell.tables:
        flat = [nc.text.strip() for row in nt.rows for nc in row.cells]

        # Pattern (a): bare "N." followed by non-empty, non-label value
        for i, t in enumerate(flat):
            if re.match(r'^[1-5]\.$', t) and i + 1 < len(flat):
                next_val = flat[i + 1].strip()
                if next_val and not re.match(r'^[1-5]\.$', next_val):
                    return True

        # Pattern (b): combined "N. value" where value is not just another "M."
        for t in flat:
            m = re.match(r'^([1-5])[\.\s]+(.+)$', t)
            if m:
                value_part = m.group(2).strip()
                # Ensure value is not just another bare label like "2." or "4."
                if value_part and not re.match(r'^[1-5]\.$', value_part):
                    return True

    return False


def extract_nested_table_choices(cell):
    """
    Extract choices from nested table.
    Handles two layouts:
      a) Alternating cells: "1." | value | "2." | value | ...
      b) Single cells with combined text: "1. 276" / "2.  287" / ...
    """
    choices = {}
    for nt in cell.tables:
        all_cells = []
        for row in nt.rows:
            for nc in row.cells:
                all_cells.append(nc.text.strip())

        # Try layout (a): alternating label / value pairs
        i = 0
        while i < len(all_cells) - 1:
            label = all_cells[i]
            value = all_cells[i + 1]
            m = re.match(r'^([1-5])\.$', label)
            if m:
                choices[int(m.group(1))] = value
                i += 2
            else:
                i += 1

        # Try layout (b): single-cell "N. value"
        if not choices:
            for text in all_cells:
                m = re.match(r'^([1-5])[\.\s]+(.+)$', text)
                if m:
                    choices[int(m.group(1))] = m.group(2).strip()
    return choices


# ─── Q156 special handler ───────────────────────────────────────────────────

def extract_q156_choices(cell_elem):
    """
    Q156 choice line: "1. [√5] 2. [√22] 3. [√32] 4. 2 5. 5"
    math_blocks = ['5', '22', '32'] (sqrt arguments)
    plain text for choices 4 and 5 = "2" and "5" respectively.
    """
    math_blocks = extract_math_blocks(cell_elem)
    # math block texts are: '5' -> √5, '22' -> √22, '32' -> √32
    sqrt_choices = {
        1: f'√{math_blocks[0]}' if len(math_blocks) > 0 else '',
        2: f'√{math_blocks[1]}' if len(math_blocks) > 1 else '',
        3: f'√{math_blocks[2]}' if len(math_blocks) > 2 else '',
    }
    # Choice 4 = "2", choice 5 = "5" are plain text in the cell
    plain_choices = {4: '2', 5: '5'}
    return {**sqrt_choices, **plain_choices}


# ─── Main command ────────────────────────────────────────────────────────────

class Command(BaseCommand):
    help = 'Fix empty (and wrong) choices for 33 problem questions by re-reading docx'

    def add_arguments(self, parser):
        parser.add_argument(
            'docx_path',
            nargs='?',
            default=r'E:\Project Peyo Peyo\Project ครูวิทย์\คลังข้อสอบ.docx',
            help='Path to source .docx file',
        )
        parser.add_argument(
            '--dry-run',
            action='store_true',
            help='Show what would be updated without saving',
        )

    def handle(self, *args, **options):
        docx_path = options['docx_path']
        dry_run = options['dry_run']

        self.stdout.write(f'Loading docx: {docx_path}')
        doc = docx.Document(docx_path)
        table = doc.tables[0]
        rows_from2 = table.rows[2:]

        problem_numbers = [
            2, 3, 4, 6, 7, 8,
            41, 42, 43, 44, 46,
            52, 54, 57, 58, 59,
            63, 64, 66,
            69, 73, 76, 79, 81, 84, 87, 88,
            91, 92, 95,
            124, 125, 156,
        ]

        updated_choices = 0
        skipped_image = []
        skipped_no_data = []
        errors = []

        for db_num in problem_numbers:
            row_idx = db_num - 1
            if row_idx >= len(rows_from2):
                errors.append(f'Q{db_num}: row index out of range')
                continue

            cell = rows_from2[row_idx].cells[1]

            try:
                q = Question.objects.get(number=db_num)
            except Question.DoesNotExist:
                errors.append(f'Q{db_num}: not in DB')
                continue

            correct_answer = q.correct_answer
            math_blocks = extract_math_blocks(cell._tc)
            has_nested = len(cell.tables) > 0

            # Count images
            ns_drawing = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            drawings = (
                cell._tc.findall(f'.//{{{ns_drawing}}}inline')
                + cell._tc.findall(f'.//{{{ns_drawing}}}anchor')
            )

            # ── Classify and extract ──────────────────────────────────────────
            overwrite_all = False  # True when existing choice1/choice3 are wrong

            if db_num == 156:
                new_choices = extract_q156_choices(cell._tc)
                overwrite_all = True

            elif len(drawings) >= 5:
                # Image-based choices
                new_choices = {i: '[รูปภาพ - ดูในเอกสารต้นฉบับ]' for i in range(1, 6)}
                overwrite_all = True
                skipped_image.append(db_num)

            elif has_nested and nested_table_is_choice_table(cell):
                # Nested table contains actual choice text
                new_choices = extract_nested_table_choices(cell)
                # These questions had all-empty choices -> no overwrite needed, but safe to do
                overwrite_all = True

            elif len(math_blocks) == 5:
                # 5 OMML blocks = 5 choices, mapping 1:1
                # choice1 may contain wrong "2." and choice3 may contain "4." -> overwrite all
                new_choices = {i + 1: math_blocks[i] for i in range(5)}
                overwrite_all = True

            elif len(math_blocks) > 5:
                # Stem contains some math; choices are per-paragraph with "N. [math]"
                new_choices = extract_math_by_para(cell._tc)
                overwrite_all = True

            else:
                new_choices = {}

            if not new_choices:
                skipped_no_data.append(db_num)
                self.stdout.write(
                    self.style.WARNING(f'Q{db_num}: no choices extracted, skipping')
                )
                continue

            # ── Write to DB ──────────────────────────────────────────────────
            choices_qs = {c.number: c for c in q.choices.all()}
            q_updated = False

            for choice_num, choice_body in new_choices.items():
                if choice_num not in choices_qs:
                    continue
                c = choices_qs[choice_num]
                should_update = overwrite_all or (not c.body.strip() and choice_body.strip())
                if should_update and choice_body.strip():
                    old_val = repr(c.body[:40]) if c.body else '""'
                    if dry_run:
                        self.stdout.write(
                            f'[DRY] Q{db_num} choice{choice_num}: {old_val} -> {repr(choice_body[:60])}'
                        )
                    else:
                        c.body = choice_body
                        c.save()
                    q_updated = True
                    updated_choices += 1

            if q_updated:
                ans_text = new_choices.get(
                    int(correct_answer) if correct_answer else -1, '???'
                )
                self.stdout.write(
                    self.style.SUCCESS(
                        f'Q{db_num} [ans={correct_answer}]: updated. '
                        f'correct choice={repr(str(ans_text)[:60])}'
                    )
                )
            else:
                self.stdout.write(f'Q{db_num}: no changes needed')

        # ── Summary ──────────────────────────────────────────────────────────
        self.stdout.write('')
        self.stdout.write(self.style.SUCCESS(f'Done: {updated_choices} choice fields updated'))
        if skipped_image:
            self.stdout.write(
                self.style.WARNING(f'Image placeholder inserted: {skipped_image}')
            )
        if skipped_no_data:
            self.stdout.write(self.style.ERROR(f'No data found: {skipped_no_data}'))
        for e in errors:
            self.stdout.write(self.style.ERROR(e))

        # ── Verification ─────────────────────────────────────────────────────
        self.stdout.write('')
        self.stdout.write('Verification:')
        remaining = []
        for q in Question.objects.prefetch_related('choices', 'parameter').order_by('number'):
            ans = q.correct_answer
            if ans:
                try:
                    c = q.choices.filter(number=int(ans)).first()
                    if c and not c.body.strip():
                        remaining.append(q.number)
                except (ValueError, TypeError):
                    pass
        if remaining:
            self.stdout.write(
                self.style.ERROR(
                    f'Remaining empty-answer issues: {len(remaining)} ข้อ: {remaining}'
                )
            )
        else:
            self.stdout.write(
                self.style.SUCCESS('All correct-answer choices are non-empty!')
            )
