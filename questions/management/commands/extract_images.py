from django.core.management.base import BaseCommand
from django.core.files.base import ContentFile
from questions.models import Question
import docx
from docx.oxml.ns import qn


class Command(BaseCommand):
    help = 'Extract images from docx and update existing questions'

    def add_arguments(self, parser):
        parser.add_argument('docx_path', type=str)

    def handle(self, *args, **options):
        path = options['docx_path']
        doc = docx.Document(path)
        table = doc.tables[0]
        updated = 0
        skipped = 0

        questions = list(Question.objects.order_by('number'))
        rows = [row for row in table.rows[2:] if row.cells[1].text.strip()]

        for i, row in enumerate(rows):
            if i >= len(questions):
                break
            question = questions[i]
            cell = row.cells[1]
            image_data = self._extract_image(cell, doc)
            if image_data:
                filename = f'question_{question.number}.png'
                question.image.save(filename, ContentFile(image_data), save=True)
                updated += 1
                self.stdout.write(f'  Saved image for question {question.number}')
            else:
                skipped += 1

        self.stdout.write(self.style.SUCCESS(
            f'Done: {updated} images saved, {skipped} had no image'
        ))

    def _extract_image(self, cell, doc):
        for para in cell.paragraphs:
            for run in para.runs:
                drawings = run._element.findall('.//' + qn('a:blip'))
                for drawing in drawings:
                    rId = drawing.get(
                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                    )
                    if rId:
                        try:
                            image_part = doc.part.related_parts[rId]
                            return image_part.blob
                        except (KeyError, AttributeError):
                            pass
        return None
